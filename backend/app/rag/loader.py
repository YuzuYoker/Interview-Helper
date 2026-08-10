"""文档加载器注册表：{扩展名: 加载函数}，加载函数返回 list[Document]。

PDF 按页读取，每页一个 Document（保留页码 metadata，分块后 chunk 继承页码）；
图片走 DashScope Qwen-VL-Max API 转文字。
"""
from typing import Callable

from langchain_core.documents import Document

from app.utils.config import settings


def load_pdf(path: str, filename: str) -> list[Document]:
    from pypdf import PdfReader

    reader = PdfReader(path)
    docs = []
    for i, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            docs.append(
                Document(
                    page_content=text,
                    metadata={"source": filename, "page": i, "type": "pdf"},
                )
            )
    return docs


def load_docx(path: str, filename: str) -> list[Document]:
    from docx import Document as DocxDocument

    doc = DocxDocument(path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [
        Document(
            page_content=text,
            metadata={"source": filename, "page": 1, "type": "docx"},
        )
    ]


def load_xlsx(path: str, filename: str) -> list[Document]:
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    docs = []
    for ws in wb.worksheets:  # 每张 sheet 一个 Document，行号入内容
        lines = []
        for idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
            vals = [str(c) for c in row if c is not None]
            if vals:
                lines.append(f"第{idx}行: " + " | ".join(vals))
        if lines:
            docs.append(
                Document(
                    page_content="\n".join(lines),
                    metadata={"source": filename, "sheet": ws.title, "page": 1, "type": "xlsx"},
                )
            )
    return docs


def load_txt(path: str, filename: str) -> list[Document]:
    with open(path, encoding="utf-8") as f:
        text = f.read()
    return [
        Document(
            page_content=text,
            metadata={"source": filename, "page": 1, "type": "txt"},
        )
    ]


def load_image(path: str, filename: str) -> list[Document]:
    """图片 OCR：走 DashScope OpenAI 兼容端点调用 qwen-vl-max 转文字。"""
    import base64
    import mimetypes

    from openai import OpenAI

    if not settings.dashscope_api_key:
        raise ValueError("DASHSCOPE_API_KEY 未配置，无法识别图片")

    client = OpenAI(
        api_key=settings.dashscope_api_key,
        base_url=settings.dashscope_base_url,
    )
    mime = mimetypes.guess_type(filename)[0] or "image/jpeg"
    with open(path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode()
    resp = client.chat.completions.create(
        model=settings.vision_model,  # qwen-vl-max
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:{mime};base64,{b64}"},
                    },
                    {
                        "type": "text",
                        "text": "请完整提取图片中的全部文字信息，保留原有层级和格式，不要解释。",
                    },
                ],
            }
        ],
        max_tokens=2048,
    )
    text = resp.choices[0].message.content or ""
    return [
        Document(
            page_content=text,
            metadata={"source": filename, "page": 1, "type": "image"},
        )
    ]


# 注册表：支持的上传类型
LOADERS: dict[str, Callable[[str, str], list[Document]]] = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".xlsx": load_xlsx,
    ".txt": load_txt,
    ".jpg": load_image,
    ".jpeg": load_image,
    ".png": load_image,
    ".bmp": load_image,
    ".webp": load_image,
}
